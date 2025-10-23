"""
SERVIÇO DE EXTRAÇÃO DE TEXTO
Plataforma Jurídica Multi-Agent

CONTEXTO DE NEGÓCIO:
Este serviço é fundamental para o fluxo de ingestão de documentos jurídicos.
Documentos podem chegar em diferentes formatos (PDFs com texto selecionável,
PDFs escaneados/imagens, arquivos DOCX). Este módulo é responsável por extrair
o texto de PDFs e DOCX para que possam ser vetorizados e armazenados no RAG.

IMPORTANTE: Este módulo NÃO faz OCR (processamento de PDFs escaneados).
O OCR será implementado em um serviço separado (TAREFA-005).

RESPONSABILIDADES:
1. Extrair texto de PDFs que contêm texto selecionável (usando PyPDF2)
2. Extrair texto de arquivos DOCX (usando python-docx)
3. Detectar se um PDF é escaneado (imagem) ou contém texto
4. Fornecer metadados sobre a extração (número de páginas, confiança, etc.)

DEPENDÊNCIAS:
- PyPDF2: Para leitura de PDFs com texto
- python-docx: Para leitura de arquivos DOCX
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional

# Bibliotecas de terceiros para processamento de documentos
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None  # Será validado nas funções que usam

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # Será validado nas funções que usam


# ==========================================
# CONFIGURAÇÃO DE LOGGING
# ==========================================
# Logging detalhado é crucial para depuração, especialmente quando lidamos
# com documentos variados que podem ter formatações inesperadas

logger = logging.getLogger(__name__)


# ==========================================
# EXCEÇÕES PERSONALIZADAS
# ==========================================
# Exceções específicas facilitam o tratamento de erros pelos chamadores
# e tornam o código mais autodocumentado

class ErroDeExtracaoDeTexto(Exception):
    """
    Exceção base para erros durante extração de texto de documentos.
    
    Todas as exceções específicas deste módulo herdam desta classe,
    permitindo que código chamador capture todos os erros relacionados
    a extração com um único except.
    """
    pass


class ArquivoNaoEncontradoError(ErroDeExtracaoDeTexto):
    """
    Levantada quando o arquivo especificado não existe no caminho fornecido.
    """
    pass


class TipoDeArquivoNaoSuportadoError(ErroDeExtracaoDeTexto):
    """
    Levantada quando o tipo de arquivo não é suportado por este serviço.
    
    Tipos suportados: .pdf, .docx
    Para outros tipos (imagens PNG/JPG), ver serviço de OCR (TAREFA-005)
    """
    pass


class DependenciaNaoInstaladaError(ErroDeExtracaoDeTexto):
    """
    Levantada quando uma biblioteca necessária (PyPDF2, python-docx) não está instalada.
    """
    pass


class PDFEscaneadoError(ErroDeExtracaoDeTexto):
    """
    Levantada quando o PDF é identificado como escaneado (imagem),
    não contendo texto extraível.
    
    Neste caso, o chamador deve direcionar o documento para o serviço de OCR.
    """
    pass


# ==========================================
# FUNÇÃO: VALIDAR EXISTÊNCIA DE ARQUIVO
# ==========================================

def validar_existencia_arquivo(caminho_arquivo: str) -> None:
    """
    Valida se o arquivo existe no caminho especificado.
    
    CONTEXTO:
    Esta é uma validação básica que deve ser feita antes de qualquer tentativa
    de processamento. Melhor falhar rápido com erro claro do que tentar
    processar e falhar com erro genérico de biblioteca.
    
    Args:
        caminho_arquivo: Caminho absoluto ou relativo para o arquivo
        
    Raises:
        ArquivoNaoEncontradoError: Se o arquivo não existir
    """
    if not os.path.exists(caminho_arquivo):
        mensagem_erro = (
            f"Arquivo não encontrado no caminho especificado: {caminho_arquivo}. "
            f"Verifique se o caminho está correto e se o arquivo não foi movido/deletado."
        )
        logger.error(mensagem_erro)
        raise ArquivoNaoEncontradoError(mensagem_erro)
    
    logger.debug(f"Arquivo validado com sucesso: {caminho_arquivo}")


# ==========================================
# FUNÇÃO: VALIDAR DEPENDÊNCIA INSTALADA
# ==========================================

def validar_dependencia_instalada(biblioteca: Any, nome_biblioteca: str) -> None:
    """
    Valida se uma biblioteca Python necessária está instalada.
    
    CONTEXTO:
    PyPDF2 e python-docx são dependências opcionais durante desenvolvimento,
    mas obrigatórias em produção. Esta validação explícita garante que
    erros de ImportError sejam tratados de forma clara.
    
    Args:
        biblioteca: O objeto da biblioteca (ou None se import falhou)
        nome_biblioteca: Nome da biblioteca para mensagem de erro
        
    Raises:
        DependenciaNaoInstaladaError: Se a biblioteca não estiver disponível
    """
    if biblioteca is None:
        mensagem_erro = (
            f"Biblioteca '{nome_biblioteca}' não está instalada. "
            f"Execute: pip install {nome_biblioteca}"
        )
        logger.error(mensagem_erro)
        raise DependenciaNaoInstaladaError(mensagem_erro)


# ==========================================
# FUNÇÃO: DETECTAR SE PDF É ESCANEADO
# ==========================================

def detectar_se_pdf_e_escaneado(caminho_arquivo_pdf: str) -> bool:
    """
    Detecta se um PDF contém texto extraível ou se é escaneado (apenas imagens).
    
    CONTEXTO DE NEGÓCIO:
    Documentos jurídicos frequentemente são processos físicos digitalizados (escaneados),
    resultando em PDFs que são basicamente imagens. Estes não podem ser processados
    com PyPDF2 e precisam passar por OCR.
    
    IMPLEMENTAÇÃO:
    A detecção é feita tentando extrair texto das primeiras páginas:
    - Se conseguir extrair texto significativo (>50 caracteres) → PDF com texto
    - Se extrair apenas espaços/caracteres especiais → PDF escaneado
    
    LIMITAÇÕES:
    - PDFs mistos (algumas páginas texto, outras escaneadas) serão identificados
      como "texto" se pelo menos uma página tiver texto
    - Não diferencia texto real de "texto" que é na verdade vetores/desenhos
    
    Args:
        caminho_arquivo_pdf: Caminho absoluto para o arquivo PDF
        
    Returns:
        True se o PDF é escaneado (imagem), False se contém texto extraível
        
    Raises:
        ArquivoNaoEncontradoError: Se o arquivo não existir
        DependenciaNaoInstaladaError: Se PyPDF2 não estiver instalado
    """
    # Validações preliminares
    validar_existencia_arquivo(caminho_arquivo_pdf)
    validar_dependencia_instalada(PdfReader, "PyPDF2")
    
    logger.info(f"Detectando tipo de PDF: {caminho_arquivo_pdf}")
    
    try:
        # Abrir o PDF para leitura
        leitor_pdf = PdfReader(caminho_arquivo_pdf)
        numero_total_de_paginas = len(leitor_pdf.pages)
        
        logger.debug(f"PDF possui {numero_total_de_paginas} página(s)")
        
        # Limite de páginas a analisar (evitar processar PDFs gigantes inteiros)
        # Analisamos até as primeiras 3 páginas ou total de páginas (o que for menor)
        numero_paginas_para_analisar = min(3, numero_total_de_paginas)
        
        # Acumulador de texto extraído
        texto_total_extraido = ""
        
        # Iterar pelas primeiras páginas
        for indice_pagina in range(numero_paginas_para_analisar):
            pagina_atual = leitor_pdf.pages[indice_pagina]
            texto_da_pagina = pagina_atual.extract_text()
            
            if texto_da_pagina:
                texto_total_extraido += texto_da_pagina
            
            logger.debug(
                f"Página {indice_pagina + 1}: {len(texto_da_pagina)} caracteres extraídos"
            )
        
        # Limpar texto extraído (remover espaços em branco excessivos)
        texto_limpo = texto_total_extraido.strip()
        
        # HEURÍSTICA DE DETECÇÃO:
        # Se conseguimos extrair pelo menos 50 caracteres de texto válido,
        # consideramos que o PDF contém texto (não é escaneado)
        LIMIAR_CARACTERES_MINIMO = 50
        
        if len(texto_limpo) >= LIMIAR_CARACTERES_MINIMO:
            logger.info(
                f"PDF identificado como TEXTO (extraídos {len(texto_limpo)} caracteres)"
            )
            return False  # NÃO é escaneado
        else:
            logger.info(
                f"PDF identificado como ESCANEADO (apenas {len(texto_limpo)} caracteres extraídos)"
            )
            return True  # É escaneado
            
    except Exception as erro:
        # Captura qualquer erro do PyPDF2 e loga
        mensagem_erro = f"Erro ao tentar detectar tipo de PDF: {str(erro)}"
        logger.error(mensagem_erro)
        raise ErroDeExtracaoDeTexto(mensagem_erro)


# ==========================================
# FUNÇÃO: EXTRAIR TEXTO DE PDF
# ==========================================

def extrair_texto_de_pdf_texto(caminho_arquivo_pdf: str) -> Dict[str, Any]:
    """
    Extrai texto de um PDF que contém texto selecionável (não escaneado).
    
    CONTEXTO DE NEGÓCIO:
    PDFs jurídicos gerados digitalmente (petições criadas em Word e exportadas,
    sentenças geradas por sistemas processuais) contêm texto extraível.
    Este é o caso mais simples e eficiente de processamento.
    
    IMPLEMENTAÇÃO:
    1. Valida que o arquivo existe e PyPDF2 está disponível
    2. Detecta se o PDF é escaneado (e falha se for)
    3. Itera por todas as páginas extraindo texto
    4. Retorna texto completo + metadados
    
    IMPORTANTE:
    Se o PDF for detectado como escaneado, esta função levanta PDFEscaneadoError.
    O chamador deve então redirecionar para o serviço de OCR (TAREFA-005).
    
    Args:
        caminho_arquivo_pdf: Caminho absoluto para o arquivo PDF
        
    Returns:
        dict contendo:
        {
            "texto_extraido": str,              # Texto completo de todas as páginas
            "numero_de_paginas": int,           # Total de páginas processadas
            "metodo_extracao": str,             # "PyPDF2"
            "caminho_arquivo_original": str,    # Caminho do arquivo processado
            "tipo_documento": str,              # "pdf_texto"
            "paginas_vazias": list[int]         # Índices de páginas sem texto (0-indexed)
        }
        
    Raises:
        ArquivoNaoEncontradoError: Se o arquivo não existir
        DependenciaNaoInstaladaError: Se PyPDF2 não estiver instalado
        PDFEscaneadoError: Se o PDF for detectado como escaneado
        ErroDeExtracaoDeTexto: Para outros erros durante processamento
    """
    # Validações preliminares
    validar_existencia_arquivo(caminho_arquivo_pdf)
    validar_dependencia_instalada(PdfReader, "PyPDF2")
    
    logger.info(f"Iniciando extração de texto do PDF: {caminho_arquivo_pdf}")
    
    # Verificar se o PDF é escaneado
    # Se for, não podemos processar com PyPDF2 - precisa OCR
    if detectar_se_pdf_e_escaneado(caminho_arquivo_pdf):
        mensagem_erro = (
            f"O PDF '{caminho_arquivo_pdf}' foi detectado como escaneado (imagem). "
            f"Use o serviço de OCR para processar este documento."
        )
        logger.warning(mensagem_erro)
        raise PDFEscaneadoError(mensagem_erro)
    
    try:
        # Abrir o PDF
        leitor_pdf = PdfReader(caminho_arquivo_pdf)
        numero_total_de_paginas = len(leitor_pdf.pages)
        
        logger.info(f"Processando PDF com {numero_total_de_paginas} página(s)")
        
        # Acumuladores
        texto_completo = ""
        lista_paginas_vazias: List[int] = []
        
        # Iterar por todas as páginas
        for indice_pagina in range(numero_total_de_paginas):
            pagina_atual = leitor_pdf.pages[indice_pagina]
            texto_da_pagina = pagina_atual.extract_text()
            
            # Verificar se a página tem texto
            if texto_da_pagina and texto_da_pagina.strip():
                texto_completo += texto_da_pagina + "\n\n"
                logger.debug(
                    f"Página {indice_pagina + 1}: {len(texto_da_pagina)} caracteres extraídos"
                )
            else:
                # Página vazia - registrar para metadados
                lista_paginas_vazias.append(indice_pagina)
                logger.warning(f"Página {indice_pagina + 1}: SEM TEXTO (vazia ou escaneada)")
        
        # Limpar texto final (remover espaços excessivos)
        texto_completo = texto_completo.strip()
        
        # Montar resultado
        resultado = {
            "texto_extraido": texto_completo,
            "numero_de_paginas": numero_total_de_paginas,
            "metodo_extracao": "PyPDF2",
            "caminho_arquivo_original": caminho_arquivo_pdf,
            "tipo_documento": "pdf_texto",
            "paginas_vazias": lista_paginas_vazias
        }
        
        logger.info(
            f"Extração concluída: {len(texto_completo)} caracteres, "
            f"{len(lista_paginas_vazias)} página(s) vazia(s)"
        )
        
        return resultado
        
    except PDFEscaneadoError:
        # Re-raise - já foi tratado acima
        raise
    except Exception as erro:
        # Captura qualquer outro erro
        mensagem_erro = f"Erro ao extrair texto do PDF: {str(erro)}"
        logger.error(mensagem_erro)
        raise ErroDeExtracaoDeTexto(mensagem_erro)


# ==========================================
# FUNÇÃO: EXTRAIR TEXTO DE DOCX
# ==========================================

def extrair_texto_de_docx(caminho_arquivo_docx: str) -> Dict[str, Any]:
    """
    Extrai texto de um arquivo Microsoft Word (.docx).
    
    CONTEXTO DE NEGÓCIO:
    Muitas petições e documentos jurídicos são criados em Word.
    Clientes podem fazer upload diretamente do .docx sem converter para PDF.
    A extração de DOCX é geralmente mais confiável que PDF pois preserva
    a estrutura do documento.
    
    IMPLEMENTAÇÃO:
    1. Valida que o arquivo existe e python-docx está disponível
    2. Abre o documento usando a biblioteca python-docx
    3. Extrai texto de todos os parágrafos
    4. Extrai texto de tabelas (comum em documentos jurídicos)
    5. Retorna texto completo + metadados
    
    IMPORTANTE:
    python-docx só funciona com formato .docx (Office 2007+).
    Arquivos .doc antigos (Office 2003) não são suportados.
    
    Args:
        caminho_arquivo_docx: Caminho absoluto para o arquivo .docx
        
    Returns:
        dict contendo:
        {
            "texto_extraido": str,              # Texto completo do documento
            "numero_de_paragrafos": int,        # Total de parágrafos
            "numero_de_tabelas": int,           # Total de tabelas
            "metodo_extracao": str,             # "python-docx"
            "caminho_arquivo_original": str,    # Caminho do arquivo processado
            "tipo_documento": str               # "docx"
        }
        
    Raises:
        ArquivoNaoEncontradoError: Se o arquivo não existir
        DependenciaNaoInstaladaError: Se python-docx não estiver instalado
        TipoDeArquivoNaoSuportadoError: Se o arquivo não for .docx válido
        ErroDeExtracaoDeTexto: Para outros erros durante processamento
    """
    # Validações preliminares
    validar_existencia_arquivo(caminho_arquivo_docx)
    validar_dependencia_instalada(DocxDocument, "python-docx")
    
    # Validar extensão do arquivo
    extensao_arquivo = Path(caminho_arquivo_docx).suffix.lower()
    if extensao_arquivo != ".docx":
        mensagem_erro = (
            f"Arquivo '{caminho_arquivo_docx}' não é um .docx válido. "
            f"Extensão detectada: {extensao_arquivo}. "
            f"Apenas arquivos .docx (Office 2007+) são suportados."
        )
        logger.error(mensagem_erro)
        raise TipoDeArquivoNaoSuportadoError(mensagem_erro)
    
    logger.info(f"Iniciando extração de texto do DOCX: {caminho_arquivo_docx}")
    
    try:
        # Abrir o documento
        documento = DocxDocument(caminho_arquivo_docx)
        
        # Acumuladores
        texto_completo = ""
        numero_de_paragrafos = 0
        
        # ETAPA 1: Extrair texto dos parágrafos
        # Parágrafos são a estrutura principal de texto em Word
        logger.debug("Extraindo texto dos parágrafos...")
        for paragrafo in documento.paragraphs:
            texto_paragrafo = paragrafo.text
            if texto_paragrafo.strip():  # Ignorar parágrafos vazios
                texto_completo += texto_paragrafo + "\n"
                numero_de_paragrafos += 1
        
        logger.debug(f"Extraídos {numero_de_paragrafos} parágrafos")
        
        # ETAPA 2: Extrair texto das tabelas
        # Documentos jurídicos frequentemente têm tabelas com dados estruturados
        numero_de_tabelas = len(documento.tables)
        logger.debug(f"Processando {numero_de_tabelas} tabela(s)...")
        
        for indice_tabela, tabela in enumerate(documento.tables):
            texto_completo += f"\n[TABELA {indice_tabela + 1}]\n"
            
            # Iterar por linhas da tabela
            for linha in tabela.rows:
                # Iterar por células da linha
                textos_celulas = [celula.text.strip() for celula in linha.cells]
                # Juntar células com separador de tabulação
                linha_texto = "\t".join(textos_celulas)
                texto_completo += linha_texto + "\n"
            
            texto_completo += f"[FIM TABELA {indice_tabela + 1}]\n\n"
        
        # Limpar texto final
        texto_completo = texto_completo.strip()
        
        # Montar resultado
        resultado = {
            "texto_extraido": texto_completo,
            "numero_de_paragrafos": numero_de_paragrafos,
            "numero_de_tabelas": numero_de_tabelas,
            "metodo_extracao": "python-docx",
            "caminho_arquivo_original": caminho_arquivo_docx,
            "tipo_documento": "docx"
        }
        
        logger.info(
            f"Extração concluída: {len(texto_completo)} caracteres, "
            f"{numero_de_paragrafos} parágrafos, {numero_de_tabelas} tabelas"
        )
        
        return resultado
        
    except Exception as erro:
        # Captura qualquer erro
        mensagem_erro = f"Erro ao extrair texto do DOCX: {str(erro)}"
        logger.error(mensagem_erro)
        raise ErroDeExtracaoDeTexto(mensagem_erro)


# ==========================================
# FUNÇÃO PRINCIPAL: EXTRAIR TEXTO (ROTEADOR)
# ==========================================

def extrair_texto_de_documento(caminho_arquivo: str) -> Dict[str, Any]:
    """
    Função principal que detecta o tipo de arquivo e chama o extrator apropriado.
    
    CONTEXTO:
    Esta é a função de "fachada" (facade pattern) que deve ser usada por
    outros módulos do sistema. Ela abstrai a complexidade de decidir qual
    método de extração usar baseado na extensão do arquivo.
    
    FLUXO DE DECISÃO:
    - .pdf → extrair_texto_de_pdf_texto()
    - .docx → extrair_texto_de_docx()
    - outros → erro (imagens devem ir para OCR - TAREFA-005)
    
    Args:
        caminho_arquivo: Caminho absoluto para o arquivo (PDF ou DOCX)
        
    Returns:
        dict com texto extraído e metadados (formato varia por tipo)
        
    Raises:
        ArquivoNaoEncontradoError: Se o arquivo não existir
        TipoDeArquivoNaoSuportadoError: Se extensão não for .pdf ou .docx
        PDFEscaneadoError: Se PDF for escaneado (precisa OCR)
        ErroDeExtracaoDeTexto: Para outros erros
    """
    validar_existencia_arquivo(caminho_arquivo)
    
    # Detectar extensão do arquivo
    extensao_arquivo = Path(caminho_arquivo).suffix.lower()
    
    logger.info(
        f"Roteando extração de texto: arquivo={caminho_arquivo}, extensão={extensao_arquivo}"
    )
    
    # Roteamento baseado em extensão
    if extensao_arquivo == ".pdf":
        return extrair_texto_de_pdf_texto(caminho_arquivo)
    
    elif extensao_arquivo == ".docx":
        return extrair_texto_de_docx(caminho_arquivo)
    
    else:
        # Tipo de arquivo não suportado
        mensagem_erro = (
            f"Tipo de arquivo '{extensao_arquivo}' não suportado para extração de texto. "
            f"Tipos suportados: .pdf, .docx. "
            f"Para imagens (.png, .jpg, .jpeg), use o serviço de OCR (TAREFA-005)."
        )
        logger.error(mensagem_erro)
        raise TipoDeArquivoNaoSuportadoError(mensagem_erro)
